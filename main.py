import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
import re
from datetime import datetime

# Load .env file
load_dotenv()

client = OpenAI(
    api_key = os.getenv("OPENAI_API_KEY"),
)

with open("input.txt", "r", encoding="utf-8") as file:
    job_descrip = file.read()

completion = client.chat.completions.create(
    model="gpt-4o",
    messages=[
        {
            "role": "user",
            "content": f"""Ahmed Elshabasi
(587)-435-9647 – Calgary, AB – ahmed.elshabasi@ucalgary.ca – LinkedIn: linkedin.com/in/ahmed-elshabasi/
Skills
•	Programming languages: Java, Python, JavaScript, C, C++, HTML, CSS, React, Node.JS, Express.JS, SQL
•	Software tools: VS Code, Git, Github, Gitlab, Unity, Unreal Engine
•	Algorithm and Data Structures: Studied different algorithms and structures in university
•	Professional Skills: Adaptability, Communication, Detail-oriented, Leadership, and Time Management
Experience
Undergraduate Research Assistant (Node, React, JS)					                           May 2024 – Sep 2024
University of Calgary,     Calgary, AB
•	Developed an automated workflow using Node and React for extracting in-depth information about the data in the corpus completing tasks within set deadlines.
•	Collected first-person videos, spoken recordings, and biometric data for a corpus of information needs for outdoor running.
•	Quickly learned new tools and technologies (Node, React, and smart capture devices) to automate data analysis workflows.
Projects
Self-Checkout Machine Software (Java)								          Sep 2023 – Dec 2023
•	Collaborated with a team of 20 to designed and develop the software for a self-checkout machine
•	Focused on user-friendly interface design and efficient transaction handling to ensure smooth customer experience.
•	Integrated functionalities that simulate real-world use cases.
Educational Assessment Web App (JS, CSS, HTML)						            Jan 2024 – Apr 2024
•	Collaborated with a team of 5 to design a web application for dynamic assessments.
•	Implemented functionality to generate random questions per session, offering immediate grading and feedback.
•	Emphasized designing a user-friendly interface to provide smooth navigation and create an engaging test experience.
Full-stack Financial Assistant| Hackathon Project (Node, React, JS)	   		     	                              Feb 2024
•	Led a team of 4 to build a full-stack prompt-based financial assistant.
•	 Used ChatGPT’s API for real-time financial insights and assistance.
•	Ensured seamless deployment within the time constraints of the hackathon (24 hours).
Education
Bachelor of Computer Science									   2022 – 2026 [Expected]
University of Calgary,  Calgary, AB									    GPA: 3.68
•	Awards:
o	PURE award
o	President’s Admission Scholarship
o	University of Calgary International Undergraduate Award
•	Certificates:
o	Google Cybersecurity Professional Certificate		
o	Ready for Research Micro credential
Volunteering
Setup Crew											          Jan 2022 – May 2022
G.N.P. Hospital,    Jeddah, Saudi Arabia
•	Assisted medical students with a staff team of 5 and documented the students’ progress.
•	Collaboratively smoothed the experience for the students by getting their feedback and answering inquiries.
•	Recorded students’ progress and managed their attendance for their hospital’s academy training.
Executive Team Member									          Dec 2021 - Apr 2022
Model United Nations (MUN) at Dar Jana International School
•	Managed and participated with the MUN executive team in the organization of documents and preparation of participants
•	Prepared the hall in which the event takes place and ensured that the event proceeded as expected.
•	As one of the spokespersons (chairmen) of the event, fulfilled my duties of planning and managing the procedures of the delegates.
Extracurricular Clubs
Event Coordinator										         Sep 2023 – May 2024
Infosec Club											
•	Coordinated hands-on security labs and ethical hacking workshops, providing practical learning experiences for members.
Member											         Sep 2023 – May 2024
Competitive Programming Club									 
•	Engaged in practical programming challenges, honing my problem-solving skills through regular workshop attendance.




this is my resume and bs coverletter keep that in mind cause i am about to ask you to do stuff

Here is the job description: {job_descrip}

Write me a cover letter but i want it in a very strict format

the format should look like this:
[Date123]:[[[the actual date of today and i need you to write it as dd/mm/yyyy]]]
[Company123]:[[[the actual company name]]]
[Location123]:[[[the address of the company]]]
[Subject123]:[[[the actual subject of the cover letter]]]
[Content123]:[[[this is what follows exactly right after "dear hiring manager," and right before "sincerely, Ahmed Elshabasi"]]]
"""
        }
    ]
)

api_response = completion.choices[0].message.content

# Extract values using regex dynamically
data = {}
matches = re.findall(r"\[(.*?)\]:\[\[\[(.*?)\]\]\]", api_response, re.DOTALL)  # re.DOTALL captures multi-line content
for key, value in matches:
    data[f"[{key}]"] = value.strip()  # Strip to remove unwanted spaces/newlines

# Add today's date
today_str = datetime.now().strftime("%d/%m/%Y")
data["[Date123]"] = today_str

print(f"data: {data}")

# Load the existing Word document (template)
doc = Document("coverletter.docx")  # Make sure your template exists

# Replace placeholders in normal text (single-line values)
for para in doc.paragraphs:
    for key, value in data.items():
        if key in para.text:
            if key == "[Content123]":
                # Special handling for multi-line content
                para.clear()  # Remove placeholder text
                for line in value.split("\n"):  # Keep each paragraph separate
                    para.add_run(line)
                    para.add_run("\n")  # Add a newline for spacing
            else:
                para.text = para.text.replace(key, value)

# Save the updated document
doc.save("Updated_CoverLetter.docx")

print("Document updated successfully!")
