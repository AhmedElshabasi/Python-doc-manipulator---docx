import re
from docx import Document

# Sample API response from ChatGPT (multi-line `[Content123]` included)
api_response = """
[Date123]:[[[20/09/2024]]]
[Company123]:[[[Spartan Controls Ltd.]]]
[Location123]:[[[Calgary, AB]]]
[Subject123]:[[[Application for Co-op/Intern Student - Data Management Specialist]]]
[Content123]:[[[Dear Hiring Manager,

I am writing to express my enthusiasm for the Co-op/Intern Student - Data Management Specialist position at Spartan Controls Ltd. As a Computer Science student at the University of Calgary, I have developed a solid foundation in programming, data management, and software development, which I believe align closely with the requirements of this role.

Having worked as an Undergraduate Research Assistant, I gained experience in handling large datasets and utilizing tools like Node.js and React to automate data analysis workflows. My academic projects, such as developing a self-checkout machine software and an educational assessment web app, have honed my skills in problem-solving, team collaboration, and delivering projects under tight deadlines.

I am particularly drawn to Spartan Controls' commitment to innovative automation and customer outcomes. I am eager to contribute to the installation, data connectivity, and system integration of industrial data software applications, supporting Spartan's mission of enhancing customer safety, reliability, and productivity.

My strong communication and time management skills, coupled with my technical expertise and adaptability, make me an excellent fit for the Digital Foundations team. I am excited about the opportunity to work in a dynamic environment, collaborating on real-world projects and gaining hands-on experience with innovative digital solutions at Spartan Controls.

Thank you for considering my application. I am looking forward to the opportunity to contribute to Spartan Controls and enhance my skills. I am available for an interview at your convenience.

Sincerely,

Ahmed Elshabasi]]]
"""

# Extract values using regex dynamically
data = {}
matches = re.findall(r"\[(.*?)\]:\[\[\[(.*?)\]\]\]", api_response, re.DOTALL)  # re.DOTALL captures multi-line content
for key, value in matches:
    data[f"[{key}]"] = value.strip()  # Strip to remove unwanted spaces/newlines

print(f"data: {data}")

# Load the existing Word document (template)
doc = Document("coverletter.docx")  # Make sure your template exists

# Replace placeholders in normal text (single-line values)
for para in doc.paragraphs:
    for key, value in data.items():
        if key in para.text:
            if key == "[Content123]":
                # Special handling for multi-line content
                para.clear()  # Remove placeholder text
                for line in value.split("\n"):  # Keep each paragraph separate
                    para.add_run(line)
                    para.add_run("\n")  # Add a newline for spacing
            else:
                para.text = para.text.replace(key, value)

# Save the updated document
doc.save("Updated_CoverLetter.docx")

print("Document updated successfully!")
